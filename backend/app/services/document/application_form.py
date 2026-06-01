from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.document.base import WordTemplateBase


class ApplicationFormGenerator(WordTemplateBase):
    def __init__(
        self,
        product_name: str,
        version: str,
        doc: Document | None = None,
    ):
        super().__init__(doc)
        self.product_name = product_name
        self.version = version

    def _ensure_minimum_length(self, text: str, minimum: int, filler: str | list[str]) -> str:
        normalized = self._normalize_doc_text(text)
        if len(normalized) >= minimum:
            return normalized
        chunks = [normalized] if normalized else []
        fillers = [filler] if isinstance(filler, str) else [item for item in filler if item]
        filler_pool = [
            item
            for item in fillers
            if item not in normalized
        ] or fillers
        index = 0
        while len("".join(chunks)) < minimum:
            chunks.append(filler_pool[index % len(filler_pool)])
            index += 1
        return "".join(chunks)

    def _add_field_row(self, table, label: str, value: str) -> None:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    self._apply_run_font(run, font_name="宋体", font_size=10)
                paragraph.paragraph_format.line_spacing = 1.12

    def generate(self, profile: dict) -> Document:
        self.set_header(f"{self.product_name}{self.version} 申请表")
        self.add_page_number()
        main_functions = self._ensure_minimum_length(
            profile.get("main_functions", ""),
            500,
            [
                "系统支持统一登录、首页概览、条件检索、状态更新、结果查询和配置维护等基础能力，用于保持主要功能页面之间的连续使用体验。",
                "软件在各模块页面中集中展示关键字段、列表记录、统计摘要和操作反馈，便于围绕当前产品主题查看业务信息并完成相应处理。",
                "系统支持导出下载、结果汇总和信息留存等常用能力，使页面数据、处理结果和对外输出内容能够保持一致。",
                "通过统一的页面命名、筛选入口和状态标识，软件有助于提升日常处理效率，并便于后续查询、统计和结果核对。",
            ],
        )

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{self.product_name} 软件著作权申请表")
        self._apply_run_font(run, font_name="黑体", font_size=16, bold=True)

        table = self.doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        fields = [
            ("软件名称", profile.get("product_name", self.product_name)),
            ("版本号", profile.get("version", self.version)),
            ("软件简称", profile.get("short_name", "")),
            ("开发完成日期", profile.get("development_date", "")),
            ("软件分类", profile.get("software_category", "")),
            ("开发的硬件环境", profile.get("hardware_environment", "")),
            ("运行的硬件环境", profile.get("runtime_hardware_environment", "")),
            ("开发该软件的操作系统", profile.get("development_os", "")),
            ("源程序量", f"{profile.get('source_code_line_estimate', 0)} 行"),
            ("软件开发环境/开发工具", profile.get("development_tools", "")),
            ("该软件的运行平台/操作系统", profile.get("runtime_platform", "")),
            ("软件运行支撑环境/支持软件", profile.get("support_environment", "")),
            ("编程语言", profile.get("programming_language", "")),
            ("开发目的", profile.get("development_purpose", "")),
            ("面向领域/行业", profile.get("industry_scope", "")),
            ("软件的主要功能", main_functions),
            ("软件的技术特点", profile.get("technical_features", "")),
        ]
        for label, value in fields:
            self._add_field_row(table, label, value)
        return self.doc
