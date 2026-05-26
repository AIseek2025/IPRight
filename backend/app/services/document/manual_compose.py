from __future__ import annotations

import hashlib
import os
import re
import unicodedata
from typing import Iterable

from docx import Document

from app.services.document.base import WordTemplateBase
from app.services.document.manual import OPTIONAL_MANUAL_MODULES, REQUIRED_MANUAL_MODULES

SCREENSHOT_MARKER_RE = re.compile(r"\[\[SCREENSHOT:([^\]]+)\]\]", re.IGNORECASE)
HEADING_RE = re.compile(r"^(#{1,4})\s+(.+?)\s*$")
LIST_ITEM_RE = re.compile(r"^(\s*)([-*]|\d+\.)\s+(.+?)\s*$")
OPTIONAL_MODULE_TITLE_BY_KEY = {item["key"]: item["title"] for item in OPTIONAL_MANUAL_MODULES}


def build_variation_seed(*parts: str) -> str:
    payload = "|".join(str(part or "").strip() for part in parts if str(part or "").strip())
    if not payload:
        payload = "default"
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:16]


def _normalize_heading(text: str) -> str:
    normalized = unicodedata.normalize("NFKC", str(text or ""))
    normalized = re.sub(r"\s+", "", normalized)
    normalized = re.sub(r"[，。；：！？、】【（）“”\"'《》<>·/\\\-_]", "", normalized)
    return normalized.lower()


def _heading_aliases(title: str) -> set[str]:
    aliases = {_normalize_heading(title)}
    for part in re.split(r"[/／|]", title):
        part = part.strip()
        if part:
            aliases.add(_normalize_heading(part))
    return aliases


def _collect_markdown_headings(markdown: str) -> set[str]:
    headings: set[str] = set()
    for line in (markdown or "").splitlines():
        match = HEADING_RE.match(line.strip())
        if not match:
            continue
        headings.add(_normalize_heading(match.group(2)))
    return headings


def _heading_present(headings: set[str], title: str) -> bool:
    aliases = _heading_aliases(title)
    for alias in aliases:
        if alias in headings:
            return True
        for heading in headings:
            if alias in heading or heading in alias:
                return True
    return False


def validate_required_manual_modules(markdown: str) -> tuple[bool, list[str]]:
    headings = _collect_markdown_headings(markdown)
    missing = [module["title"] for module in REQUIRED_MANUAL_MODULES if not _heading_present(headings, module["title"])]
    return not missing, missing


def validate_optional_manual_modules(markdown: str, selected_keys: Iterable[str] | None) -> tuple[bool, list[str]]:
    if not selected_keys:
        return True, []
    headings = _collect_markdown_headings(markdown)
    missing: list[str] = []
    for key in selected_keys:
        title = OPTIONAL_MODULE_TITLE_BY_KEY.get(str(key).strip(), "")
        if title and not _heading_present(headings, title):
            missing.append(title)
    return not missing, missing


def _strip_unsupported_symbols(text: str) -> str:
    chars: list[str] = []
    for ch in text:
        if ch in {"\u200b", "\u200c", "\u200d", "\ufeff", "\ufe0f"}:
            continue
        if unicodedata.category(ch) == "So":
            continue
        chars.append(ch)
    return "".join(chars)


_COVER_SUBTITLE_VARIANTS = (
    "软件产品说明书",
    "软件操作手册",
    "产品功能与使用说明",
    "系统功能说明书",
    "软件功能操作手册",
    "软件系统说明书",
    "产品操作指南",
    "软件功能说明书",
)

_COVER_LAYOUT_VARIANTS = (
    "centered",
    "centered_with_version",
    "minimal",
    "formal",
    "compact",
)

_HEADING_NUMBERING_STYLES = (
    "none",
    "decimal",
    "chinese",
    "mixed",
    "alpha",
)


class ManualMarkdownRenderer(WordTemplateBase):
    """将 LLM 生成的 Markdown 说明书渲染为 docx，仅做最小必要排版。"""

    def __init__(
        self,
        *,
        product_name: str,
        version: str,
        screenshots_meta: list[dict] | None = None,
        arch_diagram_path: str = "",
        variation_seed: str = "",
        doc: Document | None = None,
    ):
        super().__init__(doc)
        self.product_name = product_name
        self.version = version
        self.screenshots_meta = screenshots_meta or []
        self.arch_diagram_path = arch_diagram_path
        self.variation_seed = variation_seed or "default"
        seed_digest = hashlib.sha256(self.variation_seed.encode("utf-8")).hexdigest()
        self._cover_subtitle = _COVER_SUBTITLE_VARIANTS[
            int(seed_digest[:4], 16) % len(_COVER_SUBTITLE_VARIANTS)
        ]
        self._cover_layout = _COVER_LAYOUT_VARIANTS[
            int(seed_digest[4:6], 16) % len(_COVER_LAYOUT_VARIANTS)
        ]
        self._heading_style_offset = int(seed_digest[6:8], 16) % 3
        self._heading_numbering = _HEADING_NUMBERING_STYLES[
            int(seed_digest[8:10], 16) % len(_HEADING_NUMBERING_STYLES)
        ]
        self._section_spacing_variant = int(seed_digest[10:12], 16) % 3
        self._heading_counter = 0
        self._sub_heading_counter = 0
        self._last_heading_level = 0
        self._screenshot_by_title = {
            str(item.get("page_title", "")).strip(): item
            for item in self.screenshots_meta
            if str(item.get("page_title", "")).strip()
        }
        self._screenshot_by_route = {
            str(item.get("route", "")).strip(): item
            for item in self.screenshots_meta
            if str(item.get("route", "")).strip()
        }

    def _resolve_screenshot(self, marker: str) -> dict | None:
        key = str(marker or "").strip()
        if not key:
            return None
        if key in self._screenshot_by_title:
            return self._screenshot_by_title[key]
        if key in self._screenshot_by_route:
            return self._screenshot_by_route[key]
        for item in self.screenshots_meta:
            title = str(item.get("page_title", "")).strip()
            route = str(item.get("route", "")).strip()
            if key in title or key in route:
                return item
        return None

    def _render_screenshot_marker(self, marker: str) -> None:
        meta = self._resolve_screenshot(marker)
        if not meta:
            self.add_paragraph(f"（待补充截图：{marker}）")
            return
        image_path = str(meta.get("image_path", "")).strip()
        caption = str(meta.get("caption", "")).strip() or f"图 {meta.get('page_title', marker)}"
        if image_path and os.path.exists(image_path):
            self.add_image(image_path, width_inches=6.2, max_height_inches=6.2)
            self.add_caption(caption)
        else:
            self.add_paragraph(f"（截图文件缺失：{meta.get('page_title', marker)}）")

    def _render_architecture_marker(self) -> None:
        if self.arch_diagram_path and os.path.exists(self.arch_diagram_path):
            self.add_image(self.arch_diagram_path, width_inches=6.0)
            self.add_caption(f"图1：{self.product_name}系统架构图")
        else:
            self.add_paragraph("系统架构图生成失败，请检查中文字体与图片生成链路。")

    def _render_line_with_markers(self, text: str) -> None:
        cursor = 0
        for match in SCREENSHOT_MARKER_RE.finditer(text):
            prefix = text[cursor:match.start()].strip()
            if prefix:
                self.add_paragraph(_strip_unsupported_symbols(prefix))
            marker = match.group(1).strip()
            if marker.upper() in {"ARCH", "ARCHITECTURE", "SYSTEM_ARCHITECTURE", "系统架构图"}:
                self._render_architecture_marker()
            else:
                self._render_screenshot_marker(marker)
            cursor = match.end()
        suffix = text[cursor:].strip()
        if suffix:
            self.add_paragraph(_strip_unsupported_symbols(suffix))

    def _render_cover(self, title: str) -> None:
        """根据 cover layout variant 渲染不同的封面版式。"""
        layout = self._cover_layout
        if layout == "centered":
            for _ in range(3):
                self.doc.add_paragraph()
            self.add_title(title or self.product_name, level=0)
            version_line = self.doc.add_paragraph()
            version_line.alignment = 1
            version_run = version_line.add_run(f"版本号: {self.version}")
            self._apply_run_font(version_run, font_name="宋体", font_size=14)
            subtitle = self.doc.add_paragraph()
            subtitle.alignment = 1
            subtitle_run = subtitle.add_run(self._cover_subtitle)
            self._apply_run_font(subtitle_run, font_name="宋体", font_size=14)
        elif layout == "centered_with_version":
            for _ in range(2):
                self.doc.add_paragraph()
            self.add_title(title or self.product_name, level=0)
            self.doc.add_paragraph()
            subtitle = self.doc.add_paragraph()
            subtitle.alignment = 1
            subtitle_run = subtitle.add_run(self._cover_subtitle)
            self._apply_run_font(subtitle_run, font_name="宋体", font_size=16)
            self.doc.add_paragraph()
            version_line = self.doc.add_paragraph()
            version_line.alignment = 1
            version_run = version_line.add_run(f"{self.version}")
            self._apply_run_font(version_run, font_name="宋体", font_size=12)
        elif layout == "minimal":
            for _ in range(4):
                self.doc.add_paragraph()
            self.add_title(title or self.product_name, level=0)
            info_line = self.doc.add_paragraph()
            info_line.alignment = 1
            info_run = info_line.add_run(f"{self._cover_subtitle}  {self.version}")
            self._apply_run_font(info_run, font_name="宋体", font_size=12)
        elif layout == "formal":
            for _ in range(2):
                self.doc.add_paragraph()
            subtitle_top = self.doc.add_paragraph()
            subtitle_top.alignment = 1
            subtitle_top_run = subtitle_top.add_run(self._cover_subtitle)
            self._apply_run_font(subtitle_top_run, font_name="宋体", font_size=14)
            self.doc.add_paragraph()
            self.add_title(title or self.product_name, level=0)
            self.doc.add_paragraph()
            version_line = self.doc.add_paragraph()
            version_line.alignment = 1
            version_run = version_line.add_run(f"版本号: {self.version}")
            self._apply_run_font(version_run, font_name="宋体", font_size=12)
        else:
            for _ in range(3):
                self.doc.add_paragraph()
            self.add_title(title or self.product_name, level=0)
            version_line = self.doc.add_paragraph()
            version_line.alignment = 1
            version_run = version_line.add_run(f"{self.version} · {self._cover_subtitle}")
            self._apply_run_font(version_run, font_name="宋体", font_size=13)
        self.doc.add_page_break()

    def _format_heading_with_numbering(self, title: str, level: int) -> str:
        """根据 heading numbering style 为标题添加编号前缀。"""
        style = self._heading_numbering
        if level != self._last_heading_level:
            if level == 2:
                self._heading_counter += 1
                self._sub_heading_counter = 0
            elif level == 3:
                self._sub_heading_counter += 1
            self._last_heading_level = level
        elif level == 2:
            self._heading_counter += 1
            self._sub_heading_counter = 0
        elif level == 3:
            self._sub_heading_counter += 1

        if style == "none":
            return title
        elif style == "decimal":
            if level == 2:
                return f"{self._heading_counter}. {title}"
            elif level == 3:
                return f"{self._heading_counter}.{self._sub_heading_counter} {title}"
            return title
        elif style == "chinese":
            cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
                        "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十"]
            if level == 2:
                idx = min(self._heading_counter - 1, len(cn_nums) - 1)
                return f"{cn_nums[max(idx, 0)]}、{title}"
            elif level == 3:
                return f"（{self._sub_heading_counter}）{title}"
            return title
        elif style == "mixed":
            cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
                        "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十"]
            if level == 2:
                idx = min(self._heading_counter - 1, len(cn_nums) - 1)
                return f"第{cn_nums[max(idx, 0)]}章 {title}"
            elif level == 3:
                return f"{self._heading_counter}.{self._sub_heading_counter} {title}"
            return title
        elif style == "alpha":
            if level == 2:
                letter = chr(ord('A') + min(self._heading_counter - 1, 25))
                return f"{letter}. {title}"
            elif level == 3:
                return f"{self._heading_counter}.{self._sub_heading_counter} {title}"
            return title
        return title

    def render(self, markdown: str) -> Document:
        self.set_header(f"{self.product_name}{self.version} 说明书")
        self.add_page_number()

        lines = (markdown or "").splitlines()
        paragraph_buffer: list[str] = []
        cover_done = False

        def flush_paragraph() -> None:
            if not paragraph_buffer:
                return
            text = _strip_unsupported_symbols("\n".join(paragraph_buffer).strip())
            paragraph_buffer.clear()
            if not text:
                return
            if SCREENSHOT_MARKER_RE.search(text):
                self._render_line_with_markers(text)
            else:
                self.add_paragraph(text)

        for raw_line in lines:
            line = raw_line.rstrip()
            stripped = line.strip()
            if not stripped:
                flush_paragraph()
                continue

            heading_match = HEADING_RE.match(stripped)
            if heading_match:
                flush_paragraph()
                level = len(heading_match.group(1))
                title = _strip_unsupported_symbols(heading_match.group(2))
                if level == 1 and not cover_done:
                    cover_done = True
                    self._render_cover(title)
                    continue
                mapped_level = min(max(level + self._heading_style_offset, 1), 3)
                title = self._format_heading_with_numbering(title, mapped_level)
                self.add_title(title, level=mapped_level)
                continue

            list_match = LIST_ITEM_RE.match(line)
            if list_match:
                flush_paragraph()
                self.add_paragraph(f"- {_strip_unsupported_symbols(list_match.group(3))}")
                continue

            if SCREENSHOT_MARKER_RE.fullmatch(stripped):
                flush_paragraph()
                marker = SCREENSHOT_MARKER_RE.fullmatch(stripped).group(1).strip()
                if marker.upper() in {"ARCH", "ARCHITECTURE", "SYSTEM_ARCHITECTURE", "系统架构图"}:
                    self._render_architecture_marker()
                else:
                    self._render_screenshot_marker(marker)
                continue

            paragraph_buffer.append(stripped)

        flush_paragraph()

        if not cover_done:
            self._render_cover(self.product_name)

        return self.doc


def render_manual_markdown_to_docx(
    *,
    markdown: str,
    product_name: str,
    version: str,
    screenshots_meta: list[dict] | None = None,
    arch_diagram_path: str = "",
    variation_seed: str = "",
) -> Document:
    renderer = ManualMarkdownRenderer(
        product_name=product_name,
        version=version,
        screenshots_meta=screenshots_meta,
        arch_diagram_path=arch_diagram_path,
        variation_seed=variation_seed,
    )
    return renderer.render(markdown)


def build_fallback_manual_markdown(
    *,
    product_name: str,
    version: str,
    profile: dict,
    prd_summary: dict | None,
    screenshots_meta: list[dict],
    variation_seed: str,
) -> str:
    """LLM 失败时的简化 Markdown 兜底：章节顺序与标题措辞随 variation_seed 变化。"""
    modules = [str(module.get("title", "")).strip() for module in profile.get("modules", []) if str(module.get("title", "")).strip()]
    roles = profile.get("user_roles") or (prd_summary or {}).get("user_roles") or ["管理员", "业务主管", "业务专员"]
    optional_keys = list(profile.get("selected_optional_modules") or [])
    if len(optional_keys) < 4:
        pool = [item["key"] for item in OPTIONAL_MANUAL_MODULES]
        start = int(variation_seed[:4], 16) % max(len(pool) - 6, 1)
        optional_keys = pool[start : start + 4]

    section_orders = [
        ["文档说明", "引言", "软件概述", "开发设计 / 系统设计", "开发运行环境 / 软件适配环境", "功能结构说明", "角色权限说明", "业务流程说明"],
        ["引言", "文档说明", "软件概述", "开发运行环境 / 软件适配环境", "开发设计 / 系统设计", "功能结构说明", "业务流程说明", "角色权限说明"],
        ["文档说明", "引言", "开发设计 / 系统设计", "软件概述", "功能结构说明", "角色权限说明", "开发运行环境 / 软件适配环境", "业务流程说明"],
    ]
    ordered_required = section_orders[int(variation_seed[4:6], 16) % len(section_orders)]
    required_titles = {item["title"] for item in REQUIRED_MANUAL_MODULES}
    for title in required_titles:
        if title not in ordered_required and title not in {"软件使用说明", "技术特点说明"}:
            ordered_required.append(title)

    lines: list[str] = [f"# {product_name}", ""]
    for section in ordered_required:
        lines.append(f"## {section}")
        if section == "引言":
            lines.extend([
                "### 建设背景",
                f"{product_name}面向{profile.get('industry_scope', profile.get('scene', '当前行业'))}场景，围绕{profile.get('topic_label', product_name)}主题开展建设。",
                "### 建设目标",
                profile.get(
                    "development_purpose",
                    f"本软件通过统一入口与模块化页面，支撑{'、'.join(roles[:3])}等角色开展日常业务处理。",
                ),
            ])
        elif section == "开发设计 / 系统设计":
            lines.extend([
                profile.get(
                    "system_architecture_summary",
                    f"{product_name}采用前后端分层结构，由页面展示层、业务处理层与数据存储层组成。",
                ),
                "[[SCREENSHOT:ARCH]]",
            ])
        elif section == "功能结构说明":
            for module_title in modules[:8]:
                lines.append(f"### {module_title}")
                lines.append(
                    f"{module_title}模块承接{profile.get('keyword', product_name)}相关业务处理，提供查询、录入、复核与结果输出能力。"
                )
        elif section == "软件使用说明":
            continue
        else:
            lines.append(
                f"{product_name}{version}在{section}维度围绕当前任务画像组织正文，variation={variation_seed}。"
            )
        lines.append("")

    for key in optional_keys[:6]:
        title = OPTIONAL_MODULE_TITLE_BY_KEY.get(key, "")
        if title:
            lines.extend([f"## {title}", f"本章节结合{profile.get('scene', '当前业务')}场景，对{title}进行项目专属说明。", ""])

    lines.extend(["## 软件使用说明", ""])
    for index, meta in enumerate(screenshots_meta[:12], start=1):
        page_title = str(meta.get("page_title", f"页面{index}")).strip()
        lines.extend([
            f"### {page_title}",
            f"[[SCREENSHOT:{page_title}]]",
            f"{page_title}页面用于支撑{profile.get('keyword', product_name)}相关业务的查询、处理与结果确认。",
            "",
        ])

    lines.extend([
        "## 技术特点说明",
        profile.get(
            "technical_features",
            f"{product_name}采用浏览器访问模式，强调模块化页面、角色协同与结果导出能力。",
        ),
        "",
    ])

    return "\n".join(lines)
