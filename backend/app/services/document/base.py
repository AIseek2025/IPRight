from __future__ import annotations

import hashlib
import re

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image


_WORD_STYLE_PROFILES = [
    {
        "name": "song_standard",
        "body_font": "宋体",
        "title_font": "黑体",
        "cover_font": "宋体",
        "body_size": 10.5,
        "body_line_spacing": 1.18,
        "title_level_0_size": 18,
        "title_level_1_size": 15,
        "title_level_2_size": 12,
        "first_line_indent_cm": 0.74,
        "paragraph_after_pt": 2,
        "image_width_inches": 6.2,
        "caption_font_size": 9,
        "header_footer_font_size": 9,
    },
    {
        "name": "fangsong_formal",
        "body_font": "仿宋_GB2312",
        "title_font": "黑体",
        "cover_font": "仿宋_GB2312",
        "body_size": 11,
        "body_line_spacing": 1.2,
        "title_level_0_size": 18,
        "title_level_1_size": 15,
        "title_level_2_size": 12,
        "first_line_indent_cm": 0.8,
        "paragraph_after_pt": 2,
        "image_width_inches": 6.0,
        "caption_font_size": 9,
        "header_footer_font_size": 9,
    },
    {
        "name": "kaiti_readable",
        "body_font": "楷体_GB2312",
        "title_font": "黑体",
        "cover_font": "楷体_GB2312",
        "body_size": 11,
        "body_line_spacing": 1.22,
        "title_level_0_size": 17,
        "title_level_1_size": 15,
        "title_level_2_size": 12,
        "first_line_indent_cm": 0.78,
        "paragraph_after_pt": 2,
        "image_width_inches": 6.0,
        "caption_font_size": 9,
        "header_footer_font_size": 9,
    },
    {
        "name": "dengxian_compact",
        "body_font": "等线",
        "title_font": "黑体",
        "cover_font": "等线",
        "body_size": 10.5,
        "body_line_spacing": 1.16,
        "title_level_0_size": 18,
        "title_level_1_size": 15,
        "title_level_2_size": 12,
        "first_line_indent_cm": 0.72,
        "paragraph_after_pt": 1,
        "image_width_inches": 6.25,
        "caption_font_size": 9,
        "header_footer_font_size": 9,
    },
    {
        "name": "yahei_clear",
        "body_font": "微软雅黑",
        "title_font": "黑体",
        "cover_font": "微软雅黑",
        "body_size": 10.5,
        "body_line_spacing": 1.18,
        "title_level_0_size": 18,
        "title_level_1_size": 15,
        "title_level_2_size": 12,
        "first_line_indent_cm": 0.72,
        "paragraph_after_pt": 2,
        "image_width_inches": 6.15,
        "caption_font_size": 9,
        "header_footer_font_size": 9,
    },
    {
        "name": "song_dense",
        "body_font": "宋体",
        "title_font": "等线",
        "cover_font": "宋体",
        "body_size": 10,
        "body_line_spacing": 1.12,
        "title_level_0_size": 17,
        "title_level_1_size": 14,
        "title_level_2_size": 11,
        "first_line_indent_cm": 0.72,
        "paragraph_after_pt": 1,
        "image_width_inches": 6.3,
        "caption_font_size": 8.5,
        "header_footer_font_size": 8.5,
    },
    {
        "name": "fangsong_spacious",
        "body_font": "仿宋_GB2312",
        "title_font": "华文中宋",
        "cover_font": "华文中宋",
        "body_size": 11,
        "body_line_spacing": 1.24,
        "title_level_0_size": 18,
        "title_level_1_size": 15,
        "title_level_2_size": 12,
        "first_line_indent_cm": 0.82,
        "paragraph_after_pt": 3,
        "image_width_inches": 5.9,
        "caption_font_size": 9,
        "header_footer_font_size": 9,
    },
    {
        "name": "heiti_modern",
        "body_font": "等线",
        "title_font": "黑体",
        "cover_font": "黑体",
        "body_size": 10.5,
        "body_line_spacing": 1.16,
        "title_level_0_size": 18,
        "title_level_1_size": 15,
        "title_level_2_size": 12,
        "first_line_indent_cm": 0.7,
        "paragraph_after_pt": 2,
        "image_width_inches": 6.2,
        "caption_font_size": 9,
        "header_footer_font_size": 9,
    },
    {
        "name": "yahei_spacious",
        "body_font": "微软雅黑",
        "title_font": "华文中宋",
        "cover_font": "微软雅黑",
        "body_size": 11,
        "body_line_spacing": 1.24,
        "title_level_0_size": 18,
        "title_level_1_size": 15,
        "title_level_2_size": 12,
        "first_line_indent_cm": 0.8,
        "paragraph_after_pt": 3,
        "image_width_inches": 5.95,
        "caption_font_size": 9,
        "header_footer_font_size": 9,
    },
    {
        "name": "heiti_compact",
        "body_font": "黑体",
        "title_font": "等线",
        "cover_font": "黑体",
        "body_size": 10.5,
        "body_line_spacing": 1.12,
        "title_level_0_size": 17,
        "title_level_1_size": 14,
        "title_level_2_size": 11,
        "first_line_indent_cm": 0.68,
        "paragraph_after_pt": 1,
        "image_width_inches": 6.25,
        "caption_font_size": 8.5,
        "header_footer_font_size": 8.5,
    },
]


def pick_word_style_profile(seed: str | None = None) -> dict:
    if not seed:
        return dict(_WORD_STYLE_PROFILES[0])
    digest = hashlib.sha256(seed.encode("utf-8")).hexdigest()
    index = int(digest[:8], 16) % len(_WORD_STYLE_PROFILES)
    return dict(_WORD_STYLE_PROFILES[index])


class WordTemplateBase:
    def __init__(self, doc: Document | None = None, style_profile: dict | None = None):
        self.doc = doc or Document()
        self.style_profile = dict(style_profile or _WORD_STYLE_PROFILES[0])
        self._setup_page()

    def _setup_page(self) -> None:
        for section in self.doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(1.6)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        self._setup_styles()

    def _setup_styles(self) -> None:
        normal = self.doc.styles["Normal"]
        body_font = self.style_profile["body_font"]
        normal.font.name = body_font
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
        normal.font.size = Pt(self.style_profile["body_size"])

    def _apply_run_font(
        self,
        run,
        *,
        font_name: str = "宋体",
        font_size: int | float = 10.5,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic

    def _normalize_doc_text(self, text: str) -> str:
        text = self._sanitize_xml_text(text)
        text = text.replace("\xa0", " ")
        text = re.sub(r"\s+", " ", text).strip()
        text = re.sub(r"([\u4e00-\u9fff])\s+([\u4e00-\u9fff])", r"\1\2", text)
        text = re.sub(r"([\u4e00-\u9fff])\s+(V\d)", r"\1\2", text)
        text = re.sub(r"(V\d(?:\.\d+)*)\s+([\u4e00-\u9fff])", r"\1\2", text)
        text = re.sub(r"([，。；：！？、】【（）])\s+([\u4e00-\u9fff])", r"\1\2", text)
        text = re.sub(r"([\u4e00-\u9fff])\s+([，。；：！？、】【（）])", r"\1\2", text)
        return text

    def _sanitize_xml_text(self, text: str) -> str:
        if text is None:
            return ""
        text = str(text)
        cleaned: list[str] = []
        for ch in text:
            code = ord(ch)
            if (
                code == 0x9
                or code == 0xA
                or code == 0xD
                or 0x20 <= code <= 0xD7FF
                or 0xE000 <= code <= 0xFFFD
                or 0x10000 <= code <= 0x10FFFF
            ):
                cleaned.append(ch)
            else:
                cleaned.append("\uFFFD")
        return "".join(cleaned)

    def set_header(self, text: str) -> None:
        text = self._normalize_doc_text(text)
        for section in self.doc.sections:
            header = section.header
            if not header.paragraphs:
                header.add_paragraph()
            header.paragraphs[0].text = text
            header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in header.paragraphs[0].runs:
                self._apply_run_font(
                    run,
                    font_name=self.style_profile["body_font"],
                    font_size=self.style_profile["header_footer_font_size"],
                )

    def add_page_number(self) -> None:
        for section in self.doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            if footer.paragraphs:
                footer.paragraphs[0].clear()
                p = footer.paragraphs[0]
            else:
                p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            prefix = p.add_run("第 ")
            self._apply_run_font(
                prefix,
                font_name=self.style_profile["body_font"],
                font_size=self.style_profile["header_footer_font_size"],
            )
            run = p.add_run()
            fld_char1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._r.append(fld_char1)
            instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
            run._r.append(instr)
            fld_char2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run._r.append(fld_char2)
            self._apply_run_font(
                run,
                font_name=self.style_profile["body_font"],
                font_size=self.style_profile["header_footer_font_size"],
            )
            suffix = p.add_run(" 页")
            self._apply_run_font(
                suffix,
                font_name=self.style_profile["body_font"],
                font_size=self.style_profile["header_footer_font_size"],
            )

    def add_title(self, text: str, level: int = 0) -> None:
        text = self._normalize_doc_text(text)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        if level == 0:
            self._apply_run_font(
                run,
                font_name=self.style_profile["title_font"],
                font_size=self.style_profile["title_level_0_size"],
                bold=True,
            )
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(12)
        elif level == 1:
            self._apply_run_font(
                run,
                font_name=self.style_profile["title_font"],
                font_size=self.style_profile["title_level_1_size"],
                bold=True,
            )
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        else:
            self._apply_run_font(
                run,
                font_name=self.style_profile["title_font"],
                font_size=self.style_profile["title_level_2_size"],
                bold=True,
            )
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.1

    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        font_size: int | float | None = None,
        font_name: str | None = None,
    ) -> None:
        text = self._normalize_doc_text(text)
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        self._apply_run_font(
            run,
            font_name=font_name or self.style_profile["body_font"],
            font_size=font_size or self.style_profile["body_size"],
            bold=bold,
        )
        pf = p.paragraph_format
        pf.first_line_indent = Cm(self.style_profile["first_line_indent_cm"])
        pf.space_before = Pt(0)
        pf.space_after = Pt(self.style_profile["paragraph_after_pt"])
        pf.line_spacing = self.style_profile["body_line_spacing"]
        return p

    def add_image(
        self,
        image_path: str,
        width_inches: float | None = None,
        max_height_inches: float | None = None,
    ) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        width_inches = width_inches or float(self.style_profile["image_width_inches"])
        if max_height_inches:
            with Image.open(image_path) as img:
                width_px, height_px = img.size
            if width_px > 0 and height_px > 0:
                aspect_ratio = height_px / width_px
                target_height = width_inches * aspect_ratio
                if target_height > max_height_inches:
                    width_inches = max_height_inches / aspect_ratio
        run.add_picture(image_path, width=Inches(width_inches))
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    def add_caption(self, text: str) -> None:
        text = self._normalize_doc_text(text)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        self._apply_run_font(
            run,
            font_name=self.style_profile["body_font"],
            font_size=self.style_profile["caption_font_size"],
            italic=False,
        )
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0

    def add_code_block(self, code_text: str, font_size: int | float = 10) -> None:
        p = self.doc.add_paragraph()
        run = p.add_run(self._sanitize_xml_text(code_text))
        self._apply_run_font(run, font_name="Consolas", font_size=font_size)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Cm(0)
        pf.line_spacing = Pt(8.6)

    def save(self, path: str) -> None:
        self.doc.save(path)
