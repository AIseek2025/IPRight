from __future__ import annotations

import hashlib
import os
import re
import unicodedata

from docx import Document

from app.services.document.base import WordTemplateBase, pick_word_style_profile

REQUIRED_MANUAL_MODULES = [
    {"key": "introduction", "title": "引言"},
    {"key": "overview", "title": "软件产品说明"},
    {"key": "system_design", "title": "系统组成说明"},
    {"key": "runtime_environment", "title": "开发运行环境 / 软件适配环境"},
    {"key": "tech_features", "title": "软件特点说明"},
    {"key": "function_structure", "title": "功能说明"},
    {"key": "page_instructions", "title": "操作手册"},
]

OPTIONAL_MANUAL_MODULES = [
    {"key": "business_flows", "title": "业务流程说明", "description": "说明软件中的主要处理流程与页面衔接关系。"},
    {"key": "data_and_output", "title": "数据与结果说明", "description": "围绕页面数据内容、结果信息和导出内容做事实性说明。"},
    {"key": "security_and_maintenance", "title": "运行维护说明", "description": "说明访问控制、运行检查与维护要点。"},
    {"key": "version_evolution_and_change_management", "title": "版本信息说明", "description": "说明软件版本标识与版本更新信息。"},
    {"key": "appendix", "title": "附录与补充说明", "description": "补充术语与说明。"},
]


class SoftwareManualGenerator(WordTemplateBase):
    def __init__(
        self,
        product_name: str,
        version: str,
        profile: dict | None = None,
        doc: Document | None = None,
    ):
        resolved_profile = profile or {}
        super().__init__(doc, style_profile=pick_word_style_profile(self._manual_style_seed(product_name, version, resolved_profile)))
        self.product_name = product_name
        self.version = version
        self.profile = resolved_profile

    def _manual_style_seed(self, product_name: str, version: str, profile: dict) -> str:
        parts = [
            str(profile.get("design_seed") or "").strip(),
            str(profile.get("keyword") or "").strip(),
            str(profile.get("product_name") or product_name).strip(),
            str(version or "").strip(),
        ]
        return "|".join(parts)

    def _strip_unsupported_symbols(self, text: str) -> str:
        chars: list[str] = []
        for ch in text:
            if ch in {"\u200b", "\u200c", "\u200d", "\ufeff", "\ufe0f"}:
                continue
            if unicodedata.category(ch) == "So":
                continue
            chars.append(ch)
        return "".join(chars)

    def _normalize_sentence_spacing(self, text: str) -> str:
        text = text.replace("\xa0", " ")
        text = re.sub(r"\s+", " ", text).strip()
        text = re.sub(r"([\u4e00-\u9fff])\s+([\u4e00-\u9fff])", r"\1\2", text)
        text = re.sub(r"([\u4e00-\u9fff])\s+(V\d)", r"\1\2", text)
        text = re.sub(r"(V\d(?:\.\d+)*)\s+([\u4e00-\u9fff])", r"\1\2", text)
        text = re.sub(r"([，。；：！？、】【（）])\s+([\u4e00-\u9fff])", r"\1\2", text)
        text = re.sub(r"([\u4e00-\u9fff])\s+([，。；：！？、】【（）])", r"\1\2", text)
        return text

    def _sanitize_doc_text(self, text: str) -> str:
        return self._normalize_sentence_spacing(self._strip_unsupported_symbols(text))

    def _sanitize_ui_elements(self, elements: list[str]) -> list[str]:
        cleaned: list[str] = []
        seen: set[str] = set()
        for item in elements:
            normalized = self._sanitize_doc_text(item).strip(" 、")
            if not normalized or normalized in seen:
                continue
            seen.add(normalized)
            cleaned.append(normalized)
        return cleaned

    def _module_profiles(self) -> list[dict]:
        return list(self.profile.get("modules") or [])

    def _module_titles(self) -> list[str]:
        return [module.get("title", "") for module in self._module_profiles() if module.get("title")]

    def _profile_text(self, key: str, fallback: str) -> str:
        return self._sanitize_doc_text(self.profile.get(key, fallback))

    def _profile_list(self, key: str, fallback: list[str]) -> list[str]:
        values = self.profile.get(key)
        if isinstance(values, list):
            return [self._sanitize_doc_text(str(item)) for item in values if str(item).strip()]
        return fallback

    def _selected_optional_module_keys(self) -> list[str]:
        ordered_keys = [item["key"] for item in OPTIONAL_MANUAL_MODULES]
        configured = self.profile.get("selected_optional_modules")
        if isinstance(configured, list):
            valid = []
            seen: set[str] = set()
            for item in configured:
                key = str(item).strip()
                if key in ordered_keys and key not in seen:
                    valid.append(key)
                    seen.add(key)
            if valid:
                return valid
        return []

    def _module_steps(self, module: dict) -> list[str]:
        if module.get("steps"):
            return [self._sanitize_doc_text(str(step)) for step in module["steps"] if str(step).strip()]
        title = module.get("title", "当前模块")
        action = module.get("primary_action", f"处理{title}")
        return [
            f"进入{title}页面后先确认标题区、筛选区和主操作按钮，核对当前处理对象与业务范围。",
            f"通过搜索条件、列表记录和状态标签定位目标事项，并按需执行“{action}”等主操作。",
            f"完成处理后复核页面反馈、更新时间和相关记录，确认当前处理结果已经正确保存。",
        ]

    def _module_field_summary(self, module: dict) -> str:
        headers = [str(item).strip() for item in module.get("table_headers", []) if str(item).strip()]
        if not headers:
            return "页面通常由标题信息区、查询条件区、结果列表区和操作反馈区构成。"
        header_text = "页面主要展示以下信息：" + "、".join(headers[:8]) + "。"
        rows = [list(row) for row in module.get("rows", []) if isinstance(row, list)]
        if rows:
            header_text += " 页面列表区会结合实际业务记录展示对应字段值，并支持对单条记录进行查看或处理。"
        return self._sanitize_doc_text(header_text)

    def _module_business_value(self, module: dict) -> str:
        title = module.get("title", "当前模块")
        description = module.get("description", "")
        if description:
            return self._sanitize_doc_text(
                f"{description}该模块将信息采集、处理推进、结果复核与留痕沉淀集中在同一页面内，便于形成清晰稳定的业务闭环。"
            )
        return self._sanitize_doc_text(
            f"{title}模块用于承接当前任务中的关键业务步骤，能够把分散信息统一到标准页面中展示，并支持过程留痕与结果输出。"
        )

    def _module_example_record(self, module: dict) -> str:
        headers = [str(item).strip() for item in module.get("table_headers", []) if str(item).strip()]
        rows = [list(row) for row in module.get("rows", []) if isinstance(row, list)]
        if not headers or not rows:
            return "页面通常以标题区、筛选区、主表格和操作区组织一线业务数据。"
        first_row = [self._sanitize_doc_text(str(cell)) for cell in rows[0][: len(headers)]]
        pairs = [f"{header}为“{value}”" for header, value in zip(headers[:4], first_row[:4]) if value]
        if not pairs:
            return "模块数据样例会围绕编号、主题、责任角色、状态和更新时间等字段展开。"
        return self._sanitize_doc_text("例如列表首条业务记录中，" + "，".join(pairs) + "，便于用户快速理解当前模块的数据口径。")

    def _module_tech_notes(self, module: dict) -> list[str]:
        title = module.get("title", "当前模块")
        headers = [str(item).strip() for item in module.get("table_headers", []) if str(item).strip()]
        primary_action = self._sanitize_doc_text(str(module.get("primary_action", f"处理{title}")))
        focus = "、".join(headers[:4]) if headers else "标题、筛选项、结果列表与状态反馈"
        return [
            self._sanitize_doc_text(
                f"{title}页面围绕“{primary_action}”组织标题信息、查询条件、结果列表和处理入口，使相关业务记录能够集中展示。"
            ),
            self._sanitize_doc_text(
                f"模块核心字段围绕{focus}展开，便于用户在单页内完成查询、录入、复核和结果确认。"
            ),
            self._sanitize_doc_text(
                f"相关处理结果会在页面反馈信息、状态标识和记录更新时间中同步体现，便于后续查询、核对与结果留存。"
            ),
        ]

    def _core_entities_text(self) -> str:
        entities = [self._sanitize_doc_text(str(item)) for item in self.profile.get("core_entities", []) if str(item).strip()]
        if not entities:
            return "系统围绕业务资料、处理记录、状态结果和输出信息等基础数据内容组织页面展示与结果输出。"
        return self._sanitize_doc_text(
            "系统中的主要数据内容包括"
            + "、".join(entities[:6])
            + "等信息项，这些内容会在相关页面和输出结果中保持统一命名。"
        )

    def _module_data_dictionary_notes(self, module: dict) -> list[str]:
        title = module.get("title", "当前模块")
        headers = [self._sanitize_doc_text(str(item)) for item in module.get("table_headers", []) if str(item).strip()]
        if not headers:
            return [
                self._sanitize_doc_text(
                    f"{title}页面通常包含业务编号、主题对象、处理状态、更新时间和结果说明等基础字段，用于展示当前记录信息。"
                )
            ]
        notes: list[str] = []
        for idx, header in enumerate(headers[:6], start=1):
            notes.append(
                self._sanitize_doc_text(
                    f"{idx}. 字段“{header}”用于标识{title}页面中的对应信息项，便于用户在查询、列表查看和详情处理时快速识别记录内容。"
                )
            )
        return notes

    def _page_business_rule_notes(self, title: str, page_profile: dict) -> list[str]:
        primary_action = self._sanitize_doc_text(str(page_profile.get("primary_action", f"处理{title}")))
        filter_placeholder = self._sanitize_doc_text(str(page_profile.get("filter_placeholder", "")).strip())
        headers = [self._sanitize_doc_text(str(item)) for item in page_profile.get("table_headers", []) if str(item).strip()]
        field_focus = "、".join(headers[:4]) if headers else "编号、主题、状态、责任人"
        notes = [
            self._sanitize_doc_text(
                f"{title}页面的业务规则通常围绕{field_focus}等字段展开，用户在页面中进行筛选、查看和处理时，需要保证这些字段与正式业务口径完全一致。"
            ),
            self._sanitize_doc_text(
                f"页面内的主操作“{primary_action}”通常对应一个明确的业务节点，使用说明中应交代该动作的触发条件、执行结果、状态变化和后续处理关系。"
            ),
        ]
        if filter_placeholder:
            notes.append(
                self._sanitize_doc_text(
                    f"对于检索区中的“{filter_placeholder}”等条件，应明确其适用范围、默认值和组合筛选逻辑，避免用户因检索口径不一致而误判页面结果。"
                )
            )
        return notes

    def _page_collaboration_and_output_notes(self, title: str, page_profile: dict) -> list[str]:
        return [
            self._sanitize_doc_text(
                f"{title}承接当前模块中的信息录入、状态展示和结果回写，因此页面中的状态标签、更新时间、责任信息和结果摘要都应具备清晰的解释性。"
            ),
            self._sanitize_doc_text(
                f"{title}页面中的字段展示、图注说明和截图顺序需要与软件说明书正文保持一致，避免页面内容与文档表述出现偏差。"
            ),
            self._sanitize_doc_text(
                f"若该页面支持导出、汇总或结果查看，应在使用说明中交代输出内容的命名方式、字段范围与适用场景。"
            ),
        ]

    def _page_exception_notes(self, title: str, page_profile: dict) -> list[str]:
        return [
            self._sanitize_doc_text(
                f"{title}在实际运行中常见的异常包括筛选结果为空、状态回显不一致、关键字段缺失或操作反馈不明确。页面说明中应提前提示这些风险点，帮助判断当前页面是否符合预期。"
            ),
            self._sanitize_doc_text(
                f"在实施和运维过程中，{title}通常配套固定巡检项，包括页面标题、筛选入口、表格字段、主操作按钮、状态标签、时间信息和截图内容是否齐全，以便在问题出现时快速定位。"
            ),
            self._sanitize_doc_text(
                f"若页面需要作为说明书截图来源，还应同步检查页面布局稳定性、中文字体显示、数据样例完整性和图注对应关系，避免页面可运行但截图内容不可用。"
            ),
        ]

    def _variant_page_review_notes(self, title: str, page_profile: dict) -> list[str]:
        return [
            self._sanitize_doc_text(
                page_profile.get(
                    "variant_review_note",
                    f"{title}用于展示同一模块在特定筛选条件、聚焦结果或状态组合下的页面表现，便于说明业务规则在不同条件下的可视化差异。"
                )
            ),
            self._sanitize_doc_text(
                f"对于这类变体页，应重点核对筛选条件是否生效、结果数量是否合理、状态标签是否与主页面一致，以及图注描述能否准确反映当前页面所处的业务语境。"
            ),
        ]

    def _profile_focus_list(self, key: str, fallback: list[str]) -> list[str]:
        values = self.profile.get(key)
        if isinstance(values, list):
            return [self._sanitize_doc_text(str(item)) for item in values if str(item).strip()]
        return [self._sanitize_doc_text(item) for item in fallback]

    def _role_profiles(self, prd_summary: dict | None = None) -> list[str]:
        return (
            list(self.profile.get("user_roles") or [])
            or list((prd_summary or {}).get("user_roles") or [])
            or ["管理员", "业务主管", "运营专员"]
        )

    def generate_cover(self) -> None:
        for _ in range(4):
            self.doc.add_paragraph()
        self.add_title(self.product_name, level=0)
        p = self.doc.add_paragraph()
        p.alignment = 1
        run = p.add_run(f"版本号: {self.version}")
        self._apply_run_font(run, font_name=self.style_profile["cover_font"], font_size=14)
        p2 = self.doc.add_paragraph()
        p2.alignment = 1
        run2 = p2.add_run("软件说明书 / 操作手册")
        self._apply_run_font(run2, font_name=self.style_profile["cover_font"], font_size=14)
        self.doc.add_page_break()

    def generate_document_info(self, screenshots_meta: list[dict]) -> None:
        self.add_title("文档说明", level=1)
        self.add_paragraph(f"本文档为{self.product_name}{self.version}的软件说明书/操作手册。")
        self.add_paragraph(
            self._profile_text(
                "overview_version_summary",
                "文档内容主要包括引言、软件产品说明、系统组成说明、运行环境、软件特点、功能说明以及页面操作说明等内容。",
            )
        )

    def generate_introduction(self) -> None:
        self.add_title("引言", level=1)
        self.add_title("产品引言", level=2)
        self.add_paragraph(
            self._profile_text(
                "development_background",
                f"{self.product_name}围绕当前产品对应的业务场景构建，主要用于承接日常信息处理、状态查看和结果查询等软件功能。",
            )
        )
        self.add_title("文档用途", level=2)
        self.add_paragraph(
            self._profile_text(
                "development_purpose",
                "本文档用于对软件的组成、页面功能和操作方式进行客观说明，便于围绕当前软件产品形成完整、正式的产品资料。",
            )
        )

    def generate_overview(self, prd_summary: dict | None = None, modules: list[str] | None = None) -> None:
        self.add_title("软件概述", level=1)
        self.add_title("产品简介", level=2)
        module_items = modules or self._module_titles() or (prd_summary or {}).get("core_modules") or []
        self.add_paragraph(
            self._profile_text(
                "overview_product_intro",
                f"{self.product_name}是一套围绕{self.profile.get('scene', '当前产品相关业务处理与结果查看')}构建的软件产品，主要包含{'、'.join(module_items[:6]) if module_items else '软件首页与主要业务页面'}等功能内容。",
            )
        )
        self.add_paragraph(
            self._profile_text(
                "overview_version_summary",
                f"当前软件版本为{self.version}。软件以浏览器访问方式运行，并围绕页面入口、处理界面、状态反馈和结果信息组织整体功能。",
            )
        )
        self.add_title("软件主要功能", level=2)
        for mod in module_items:
            self.add_paragraph(f"- {mod}")
        self.add_title("产品目标", level=2)
        self.add_paragraph(
            self._profile_text(
                "development_purpose",
                "本软件通过统一页面和连续操作界面组织产品功能，使相关业务内容能够在同一软件中连续呈现和处理。",
            )
        )

    def generate_runtime_environment(self) -> None:
        self.add_title("开发运行环境 / 软件适配环境", level=1)
        self.add_title("运行硬件环境", level=2)
        self.add_paragraph(self.profile.get("hardware_environment", "CPU: 2核及以上；内存: 8GB及以上；磁盘: 100GB及以上。"))
        self.add_title("客户端访问环境", level=2)
        self.add_paragraph(self.profile.get("runtime_hardware_environment", "CPU: 2核及以上；内存: 4GB及以上；磁盘: 50GB及以上。"))
        self.add_title("软件环境", level=2)
        self.add_paragraph(f"运行平台/操作系统: {self.profile.get('runtime_platform', 'Linux 服务器 + 主流浏览器环境')}")
        self.add_paragraph(f"运行支撑环境/支持软件: {self.profile.get('support_environment', 'Chrome/Edge 浏览器、Node.js 18+、Python 3.11+、PostgreSQL 或 SQLite')}")
        self.add_title("适配说明", level=2)
        self.add_paragraph(
            f"{self.product_name}采用浏览器访问的软件架构，可在常见桌面浏览器环境中稳定运行，并适配日常办公场景下的常用分辨率。"
        )

    def generate_system_design(self, arch_diagram_path: str = "") -> None:
        self.add_title("系统组成说明", level=1)
        self.add_title("系统组成", level=2)
        self.add_paragraph(
            self._profile_text(
                "system_architecture_summary",
                f"{self.product_name}由登录入口、系统首页以及各业务功能页面组成，各页面围绕当前产品主题承接信息查看、状态处理和结果呈现。"
            )
        )
        self.add_paragraph(
            self._profile_text(
                "system_pipeline_summary",
                "软件通过统一页面入口、页面导航和结果反馈组织整体结构，使相关功能页面能够在同一产品界面中连续访问。"
            )
        )
        self.add_title("功能页面关系", level=2)
        module_titles = self._module_titles()
        if module_titles:
            self.add_paragraph("软件主要页面包括：" + "、".join(module_titles[:10]) + "。")
        else:
            self.add_paragraph("软件页面关系将围绕登录入口、系统首页和各业务功能页面展开。")
        self.add_title("系统组成图", level=2)
        if arch_diagram_path and os.path.exists(arch_diagram_path):
            if arch_diagram_path.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
                self.add_image(arch_diagram_path, width_inches=6.0)
                self.add_caption(f"图1：{self.product_name}系统组成图")
            else:
                self.add_paragraph("系统组成图生成失败，请检查图片生成链路。")
        else:
            self.add_paragraph("系统组成图生成失败，请检查中文字体与图片生成链路。")
        self.add_title("主要功能内容", level=2)
        self.add_paragraph(
            self._profile_text(
                "main_functions",
                "软件主要功能围绕登录访问、系统首页、业务页面处理、状态反馈和结果查看展开。",
            )
        )

    def generate_function_structure(self, modules: list[str] | None = None) -> None:
        self.add_title("功能结构说明", level=1)
        if self._module_profiles():
            for module in self._module_profiles():
                self.add_title(module["title"], level=2)
                self.add_title("功能说明", level=3)
                self.add_paragraph(module.get("description", f"{module['title']}模块用于承载该业务主题的主要操作。"))
                self.add_title("功能要点", level=3)
                for highlight in module.get("highlights", []):
                    self.add_paragraph(f"- {highlight}")
                self.add_title("页面作用", level=3)
                self.add_paragraph(self._module_business_value(module))
        else:
            modules = modules or ["软件首页", "主要业务页面"]
            for mod in modules:
                self.add_title(mod, level=2)
                self.add_paragraph(f"{mod}用于呈现当前产品中的相关功能内容和页面处理入口。")

    def generate_business_flows(self) -> None:
        self.add_title("常见业务流程说明", level=1)
        self.add_title("基础使用流程", level=2)
        self.add_paragraph(
            self._profile_text(
                "business_flow_basic",
                "使用人员首先通过登录页完成身份验证，进入系统首页后查看统计信息与待办摘要，再根据左侧导航进入具体功能模块完成录入、维护、查询、统计或配置等操作。",
            )
        )
        self.add_title("功能处理流程", level=2)
        self.add_paragraph(
            self._profile_text(
                "business_flow_materials",
                "系统围绕信息录入、结果查询、状态流转、处理反馈和结果输出形成连续处理链路，使页面操作与结果记录保持一致。",
            )
        )
        if self._module_profiles():
            self.add_title("模块协同流程", level=2)
            self.add_paragraph(
                self._profile_text(
                    "business_flow_module_collaboration",
                    f"围绕{self.product_name}的功能模块，软件支持按照模块顺序完成信息登记、处理推进、结果查看和记录留存。"
                )
            )

    def generate_data_and_output(self) -> None:
        self.add_title("数据组织与结果输出说明", level=1)
        self.add_title("数据组织说明", level=2)
        self.add_paragraph(
            self._profile_text(
                "data_organization",
                f"{self.product_name}围绕软件中的业务记录、状态信息、查询条件和结果数据组织页面内容，页面中的字段名称、状态标签、筛选项和操作入口保持统一命名方式。",
            )
        )
        self.add_title("结果输出说明", level=2)
        self.add_paragraph(
            self._profile_text(
                "result_output",
                "系统中的处理结果可通过页面回显、列表记录、详情信息和导出内容进行查看，用于保留业务处理过程与结果信息。",
            )
        )
        self.add_title("数据一致性说明", level=2)
        self.add_paragraph(
            self._profile_text(
                "material_arrangement",
                "同一类业务数据在不同页面中保持统一编号方式、字段命名和状态表达，便于对相关记录进行查询、统计和结果核对。",
            )
        )
        self.add_title("主要数据内容说明", level=2)
        self.add_paragraph(self._core_entities_text())
        self.add_paragraph(
            self._sanitize_doc_text(
                "同一类业务对象在首页统计、模块列表、详情信息和结果输出中的命名应保持统一，避免出现字段别名或状态描述不一致的情况。"
            )
        )

    def generate_appendix(self) -> None:
        self.add_title("附录与补充说明", level=1)
        self.add_title("术语与缩略语说明", level=2)
        self.add_paragraph(
            self._sanitize_doc_text(
                "本文档中的“模块”通常指具备独立页面入口和业务职责的功能单元；“工件”指任务在执行过程中生成并可复核的中间或最终产物；“导出文件”指说明书、源码文档、申请表、截图材料和下载包等输出内容。"
            )
        )
        self.add_title("维护记录说明", level=2)
        self.add_paragraph(
            self._sanitize_doc_text(
                "后续版本通常维护变更记录表，用于记录每次版本迭代涉及的模块、字段、截图、导出内容和说明书章节变化，便于持续查看版本差异。"
            )
        )
        self.add_title("模块补充清单", level=2)
        for module in self._module_profiles():
            self.add_paragraph(
                self._sanitize_doc_text(
                    f"{module.get('title', '当前模块')}：附录中补充字段口径、责任角色、主操作入口、状态流转和结果查看要点等信息，用于后续版本迭代时统一参考。"
                )
            )

    def generate_version_evolution_and_change_management(self) -> None:
        self.add_title("版本演进与变更管理说明", level=1)
        self.add_title("版本演进说明", level=2)
        self.add_paragraph(
            self._sanitize_doc_text(
                f"{self.product_name}在后续版本演进中，会围绕模块新增、字段调整、角色变化、截图更新和说明书章节同步建立变更记录，确保版本差异能够被快速理解。"
            )
        )
        self.add_paragraph(
            self._sanitize_doc_text(
                "每次版本升级时，除页面和接口本身外，还应同步复核截图清单、图注、说明书正文、源码文档目录和发布包内容，避免代码已更新而正式材料仍停留在旧版。"
            )
        )
        self.add_title("变更管理控制点", level=2)
        controls = [
            "模块标题、页面路由与角色权限变化时同步更新说明书章节",
            "字段口径、状态枚举与导出格式变化时同步更新截图和图注",
            "新增业务流程或页面变体时同步更新页面说明与截图清单",
            "发布前复核导出文件名称、版本号、页数和截图数量",
        ]
        for idx, item in enumerate(controls, start=1):
            self.add_paragraph(f"{idx}. {item}")

    def generate_security_and_maintenance(self) -> None:
        self.add_title("安全、审计与运维说明", level=1)
        self.add_title("权限与访问控制", level=2)
        self.add_paragraph(
            self._sanitize_doc_text(
                f"{self.product_name}按照账号权限控制页面访问范围，不同账号可在授权范围内查看对应数据并执行相关操作。通过统一登录、页面权限和过程记录，系统能够降低误操作风险并提升结果可追溯性。"
            )
        )
        self.add_title("日志与审计留痕", level=2)
        self.add_paragraph(
            self._sanitize_doc_text(
                "系统在任务执行、页面截图、导出生成和状态切换过程中保留关键事件记录，用于定位异常、解释结果并追溯关键节点。对于正式业务模块，也可围绕编号、状态、责任角色、更新时间等字段建立审计维度。"
            )
        )
        self.add_title("运行维护说明", level=2)
        self.add_paragraph(
            self._sanitize_doc_text(
                "运维侧需要重点关注依赖安装、前后端启动、健康检查、截图质量、工件落盘和导出文件发布等环节。若其中任一阶段异常，系统应能够通过任务时间线、阶段状态和日志信息快速定位问题，避免用户误判为长时间卡住。"
            )
        )
        self.add_paragraph(
            self._sanitize_doc_text(
                "在版本迭代过程中，模块标题、页面路由、截图数量、说明书页数和导出内容会同步更新，以确保软件功能与相关说明保持一致。"
            )
        )

    def _guess_usage_steps(self, title: str, meta: dict | None = None) -> list[str]:
        if meta and meta.get("steps"):
            return list(meta["steps"])
        return []

    def _guess_usage_description(self, title: str, elements: list[str], meta: dict | None = None) -> list[str]:
        visible = "、".join(self._sanitize_ui_elements(elements)[:10])
        description = meta.get("description") if meta else None
        primary_action = meta.get("primary_action") if meta else ""
        base = []
        if description:
            base.append(self._sanitize_doc_text(description))
        elif visible:
            base.append(self._sanitize_doc_text(f"{title}围绕{visible}等信息元素组织页面内容，用于承载当前业务主题下的关键处理动作与结果反馈。"))
        else:
            base.append(self._sanitize_doc_text(f"{title}用于承载当前业务主题下的关键处理动作、信息呈现与结果反馈。"))
        if primary_action:
            base.append(self._sanitize_doc_text(f"该页面的常用主操作为“{primary_action}”，通常位于页面主体区域或列表工具栏位置。"))
        return base

    def _is_compact_variant_page(self, meta: dict, seen_routes: set[str]) -> bool:
        title = meta.get("page_title", "")
        route = meta.get("route", "")
        scenario_id = meta.get("scenario_id", "")
        if "筛选结果" in title:
            return True
        if scenario_id.endswith("-filtered"):
            return True
        return bool(route and route in seen_routes)

    def _add_variant_page_instruction(self, title: str, caption: str, image_path: str, elements: list[str], page_profile: dict) -> None:
        self.add_title(title, level=3)
        if image_path and os.path.exists(image_path):
            self.add_image(image_path, width_inches=6.1, max_height_inches=5.8)
            self.add_caption(caption)
        self.add_paragraph(
            self._sanitize_doc_text(
                page_profile.get(
                    "variant_instruction",
                    f"{title}用于呈现在筛选、聚焦或结果定位条件下的业务处理状态，便于说明同一模块在具体业务条件下的信息组织方式、结果呈现逻辑和处理入口。",
                )
            )
        )
        if elements:
            self.add_paragraph(self._sanitize_doc_text("该功能主要涵盖：" + "、".join(elements[:10]) + "。"))
        else:
            self.add_paragraph("该功能状态下应重点关注筛选条件、结果列表、状态标签和主操作入口。")
        self.add_title("变体页核查说明", level=4)
        for note in self._variant_page_review_notes(title, page_profile):
            self.add_paragraph(note)

    def generate_page_instructions(self, screenshots_meta: list[dict]) -> None:
        self.add_title("软件操作说明", level=1)
        self.add_title("操作说明总述", level=2)
        self.add_paragraph(
            self._profile_text(
                "usage_overview",
                "系统围绕“登录 -> 首页查看 -> 进入目标功能模块 -> 完成录入、查询、统计、审核或配置操作”的基本路径组织主要功能。后续章节将按页面顺序陈述各页面的职责、信息重点与典型操作。",
            )
        )
        self.add_title("主要页面操作说明", level=2)
        if not screenshots_meta:
            self.add_paragraph("（待截图完成后自动生成页面操作说明）")
            return

        module_meta = {module["title"]: module for module in self._module_profiles()}
        seen_routes: set[str] = set()
        for i, meta in enumerate(screenshots_meta):
            title = meta.get("page_title", f"页面{i + 1}")
            route = meta.get("route", "")
            caption = meta.get("caption", "") or meta.get("suggested_caption", "") or f"图{i + 1} {title}"
            image_path = meta.get("image_path", "")
            elements = self._sanitize_ui_elements(meta.get("elements", []))
            page_profile = module_meta.get(title) or module_meta.get(title.replace("筛选结果", "").strip()) or {}
            compact_variant = self._is_compact_variant_page(meta, seen_routes)

            if route:
                seen_routes.add(route)

            if compact_variant:
                self._add_variant_page_instruction(title, caption, image_path, elements, page_profile)
                continue

            self.add_title(title, level=2)
            if image_path and os.path.exists(image_path):
                self.add_image(image_path, width_inches=6.2, max_height_inches=6.2)
                self.add_caption(caption)

            for paragraph in self._guess_usage_description(title, elements, page_profile):
                self.add_paragraph(paragraph)

            business_value = self._sanitize_doc_text(
                page_profile.get(
                    "business_value",
                    f"{title}对应当前软件中的重点业务环节，用于衔接信息录入、结果查询、状态跟踪与处理反馈。",
                )
            )
            self.add_paragraph(business_value)

            highlights = [
                self._sanitize_doc_text(str(item))
                for item in (page_profile.get("highlights") or [])
                if str(item).strip()
            ]
            if highlights:
                self.add_title("页面重点", level=3)
                for highlight in highlights[:5]:
                    self.add_paragraph(f"- {highlight}")
            elif elements:
                self.add_title("页面重点", level=3)
                self.add_paragraph(self._sanitize_doc_text("本页面重点包括：" + "、".join(elements[:12]) + "。"))

            steps = [self._sanitize_doc_text(step) for step in self._guess_usage_steps(title, page_profile) if step]
            if steps:
                self.add_title("操作流程", level=3)
                for j, step in enumerate(steps, 1):
                    self.add_paragraph(f"{j}. {step}")

    def generate_tech_features(self) -> None:
        self.add_title("软件特点说明", level=1)
        self.add_paragraph(
            self._profile_text(
                "technical_features",
                "本软件围绕统一页面入口、连续页面操作、状态反馈和结果查看等方面组织产品特点，相关功能以正式软件页面方式呈现。",
            )
        )
        self.add_paragraph(
            self._profile_text(
                "technical_feature_detail",
                f"{self.product_name}通过统一页面入口、清晰导航和连续操作路径承接各项业务处理，能够在同一软件中完成信息查看、状态处理和结果呈现。",
            )
        )
        features = self._profile_list("technical_feature_bullets", [
            "采用浏览器访问方式，软件页面可在统一入口中连续访问",
            "页面结构清晰，功能入口明确，便于围绕当前产品主题展开使用",
            "支持状态反馈、结果查看和连续页面处理",
            "支持围绕主要业务页面组织查询、处理和结果呈现",
            "页面信息与说明书内容保持对应，便于围绕产品功能形成完整资料",
        ])
        for feature in features:
            self.add_paragraph(f"- {feature}")

    def generate_full(
        self,
        prd_summary: dict | None = None,
        screenshots_meta: list[dict] | None = None,
        modules: list[str] | None = None,
        arch_diagram_path: str = "",
    ) -> Document:
        header_text = f"{self.product_name}{self.version} 说明书"
        self.set_header(header_text)
        self.add_page_number()

        screenshots_meta = screenshots_meta or []
        selected_optional_modules = set(self._selected_optional_module_keys())
        self.generate_cover()
        self.generate_document_info(screenshots_meta)
        self.generate_introduction()
        self.generate_overview(prd_summary, modules)
        self.generate_system_design(arch_diagram_path)
        self.generate_runtime_environment()
        self.generate_tech_features()
        self.generate_function_structure(modules)
        self.generate_page_instructions(screenshots_meta)
        if "business_flows" in selected_optional_modules:
            self.generate_business_flows()
        if "data_and_output" in selected_optional_modules:
            self.generate_data_and_output()
        if "security_and_maintenance" in selected_optional_modules:
            self.generate_security_and_maintenance()
        if "version_evolution_and_change_management" in selected_optional_modules:
            self.generate_version_evolution_and_change_management()
        if "appendix" in selected_optional_modules:
            self.generate_appendix()
        return self.doc
